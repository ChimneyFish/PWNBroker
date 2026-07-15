"""
Extract plain text from an uploaded policy document (.docx, .txt, .md) into the
same lightweight text convention used by policy_templates.py and policy_export.py
("# " / "## " headings, "- " bullets, "| ... |" tables), so an imported document
renders and exports the same way as a hand-written or template-generated one.
"""
from pathlib import Path

SUPPORTED_EXTS = {".docx", ".txt", ".md"}


def extract_text(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported file type: {ext}")
    if ext in (".txt", ".md"):
        return data.decode("utf-8", errors="replace")
    return _extract_docx(data)


def _extract_docx(data: bytes) -> str:
    import io
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        style = (para.style.name or "").lower() if para.style else ""
        if not text:
            lines.append("")
            continue
        if style in ("title", "heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("heading"):
            lines.append(f"## {text}")
        elif "list" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("|" + "|".join(["---"] * len(cells)) + "|")

    return "\n".join(lines).strip()
